import streamlit as st
import pandas as pd
import numpy as np
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

# Настройка страницы
st.set_page_config(
    page_title="Обработчик протоколов механических испытаний",
    page_icon="📊",
    layout="wide"
)

# Стили
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 2rem;
    }
    .info-box {
        background-color: #f0f7ff;
        padding: 1rem;
        border-radius: 10px;
        border-left: 5px solid #1E3A8A;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Нормативные значения для разных марок стали
STEEL_GRADES = {
    '20': {
        'name': 'Сталь марки 20',
        'room_temp': {
            'strength_range': (412, 549),
            'yield_min': 216,
            'elongation_min': 24,
            'reduction_min': 45
        },
        'high_temp_points': [
            (250, 196),
            (400, 137),
            (450, 127)
        ],
        'description': 'Углеродистая качественная конструкционная сталь'
    },
    '12Х1МФ': {
        'name': 'Сталь марки 12Х1МФ',
        'room_temp': {
            'strength_range': (441, 637),
            'yield_min': 274,
            'elongation_min': 21,
            'reduction_min': 55
        },
        'high_temp_points': [
            (400, 216),
            (450, 206)
        ],
        'description': 'Жаропрочная хромомолибденованадиевая сталь'
    }
}

def get_interpolated_yield(steel_grade, temp):
    """Линейная интерполяция нормативного предела текучести для выбранной марки стали"""
    if steel_grade not in STEEL_GRADES:
        return 0
    
    steel_data = STEEL_GRADES[steel_grade]
    
    # Если температура 20°C или ниже, используем комнатное значение
    if temp <= 20:
        return steel_data['room_temp']['yield_min']
    
    # Если температура выше максимальной из известных точек
    max_temp = max(t for t, _ in steel_data['high_temp_points'])
    if temp > max_temp:
        # Используем значение для максимальной температуры
        for t, value in reversed(steel_data['high_temp_points']):
            if t == max_temp:
                return value
    
    # Находим две ближайшие точки для интерполяции
    sorted_points = sorted(steel_data['high_temp_points'])
    
    # Если температура меньше минимальной известной
    min_temp = sorted_points[0][0]
    if temp < min_temp:
        # Интерполируем между комнатной температурой и первой точкой
        room_yield = steel_data['room_temp']['yield_min']
        first_temp, first_yield = sorted_points[0]
        
        if first_temp <= 20:
            return first_yield
        
        result = room_yield + (first_yield - room_yield) * (temp - 20) / (first_temp - 20)
        return round(result)
    
    # Ищем интервал для интерполяции
    for i in range(len(sorted_points) - 1):
        t1, y1 = sorted_points[i]
        t2, y2 = sorted_points[i + 1]
        
        if t1 <= temp <= t2:
            # Линейная интерполяция
            result = y1 + (y2 - y1) * (temp - t1) / (t2 - t1)
            return round(result)
    
    # Если не нашли подходящий интервал, возвращаем значение для максимальной температуры
    return sorted_points[-1][1]

def check_against_normative(value, temp, param, steel_grade, is_high_temp=False):
    """Проверка значения на соответствие нормативу"""
    if steel_grade not in STEEL_GRADES:
        return True
    
    steel_data = STEEL_GRADES[steel_grade]
    
    try:
        num_value = float(value)
    except:
        return True
    
    if temp <= 20 or not is_high_temp:
        if param == 'strength':
            min_val, max_val = steel_data['room_temp']['strength_range']
            return min_val <= num_value <= max_val
        elif param == 'yield':
            min_val = steel_data['room_temp']['yield_min']
            return num_value >= min_val
        elif param == 'elongation':
            min_val = steel_data['room_temp']['elongation_min']
            return num_value >= min_val
        elif param == 'reduction':
            min_val = steel_data['room_temp']['reduction_min']
            return num_value >= min_val
    else:
        if param == 'yield':
            normative_value = get_interpolated_yield(steel_grade, temp)
            return num_value >= normative_value
        return True
    
    return True

def clean_number(text):
    """Очистка и преобразование чисел из текста"""
    if not text:
        return 0
    
    text = str(text).replace(' ', '')
    text = text.replace(',', '.')
    text = re.sub(r'[^\d.]', '', text)
    
    try:
        return float(text) if '.' in text else int(text)
    except:
        return 0

def parse_protocol_from_docx(file_content):
    """Парсинг данных из DOCX файла с протоколом"""
    doc = Document(BytesIO(file_content))
    
    data_rows = []
    
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            
            for i, cell_text in enumerate(cells):
                if re.match(r'^\d+-\d+$', cell_text):
                    try:
                        sample_mark = cell_text
                        
                        if len(cells) >= 14:
                            temp_text = cells[5]
                            temp_match = re.search(r'(\d+)', temp_text)
                            temperature = int(temp_match.group(1)) if temp_match else 20
                            
                            strength_text = cells[10]
                            strength = clean_number(strength_text)
                            
                            yield_text = cells[11]
                            yield_strength = clean_number(yield_text)
                            
                            reduction_text = cells[12]
                            reduction = clean_number(reduction_text)
                            
                            elongation_text = cells[13]
                            elongation = clean_number(elongation_text)
                            
                            data_rows.append({
                                'Клеймо': sample_mark,
                                'Температура': temperature,
                                'Предел прочности': strength,
                                'Предел текучести': yield_strength,
                                'Отн. удл.': elongation,
                                'Отн. суж.': reduction
                            })
                            
                    except:
                        continue
    
    if not data_rows:
        return parse_protocol_from_text('\n'.join([p.text for p in doc.paragraphs]))
    
    return pd.DataFrame(data_rows)

def parse_protocol_from_text(text):
    """Парсинг данных из текста протокола"""
    lines = text.split('\n')
    data_rows = []
    
    for line in lines:
        if re.search(r'\d+-\d+', line) and any(x in line for x in ['МПа', '485', '297', '57', '30']):
            line_clean = re.sub(r'\s+', ' ', line.strip())
            parts = line_clean.split()
            
            for i, part in enumerate(parts):
                if re.match(r'^\d+-\d+$', part):
                    try:
                        sample_mark = part
                        
                        numbers = []
                        for j in range(i+1, len(parts)):
                            cleaned = clean_number(parts[j])
                            if cleaned != 0:
                                numbers.append(cleaned)
                        
                        if len(numbers) >= 12:
                            temperature = int(numbers[2]) if len(numbers) > 2 else 20
                            strength = numbers[7] if len(numbers) > 7 else 0
                            yield_strength = numbers[8] if len(numbers) > 8 else 0
                            reduction = numbers[9] if len(numbers) > 9 else 0
                            elongation = numbers[10] if len(numbers) > 10 else 0
                            
                            data_rows.append({
                                'Клеймо': sample_mark,
                                'Температура': temperature,
                                'Предел прочности': strength,
                                'Предел текучести': yield_strength,
                                'Отн. удл.': elongation,
                                'Отн. суж.': reduction
                            })
                            
                    except:
                        continue
    
    return pd.DataFrame(data_rows)

def parse_mapping_file(mapping_file):
    """Парсинг файла соответствия названий образцов"""
    try:
        if mapping_file.name.endswith('.xlsx'):
            df_mapping = pd.read_excel(mapping_file, header=None)
        else:
            return {}
        
        mapping = {}
        rows = []
        
        for idx, row in df_mapping.iterrows():
            if len(row) >= 2 and pd.notna(row[0]) and pd.notna(row[1]):
                new_name = str(row[0]).strip()
                lab_number = str(row[1]).strip()
                
                try:
                    numbers = re.findall(r'\d+', lab_number)
                    if numbers:
                        pipe_num = int(numbers[0])
                        rows.append({
                            'index': idx,
                            'pipe_num': pipe_num,
                            'new_name': new_name
                        })
                except ValueError:
                    continue
        
        rows.sort(key=lambda x: x['index'])
        
        for order, row in enumerate(rows, 1):
            mapping[row['pipe_num']] = {
                'new_name': row['new_name'],
                'order': order
            }
        
        return mapping
    except Exception as e:
        st.error(f"Ошибка при чтении файла соответствия: {str(e)}")
        return {}

def get_test_data():
    """Возвращает тестовые данные из примера протокола"""
    test_data = [
        {'Клеймо': '1-1', 'Температура': 20, 'Предел прочности': 485, 'Предел текучести': 297, 'Отн. удл.': 30, 'Отн. суж.': 57},
        {'Клеймо': '1-2', 'Температура': 20, 'Предел прочности': 481, 'Предел текучести': 295, 'Отн. удл.': 33, 'Отн. суж.': 61},
        {'Клеймо': '1-3', 'Температура': 403, 'Предел прочности': 478, 'Предел текучести': 214, 'Отн. удл.': 28, 'Отн. суж.': 63},
        {'Клеймо': '1-4', 'Температура': 403, 'Предел прочности': 483, 'Предел текучести': 289, 'Отн. удл.': 24, 'Отн. суж.': 58},
        {'Клеймо': '2-1', 'Температура': 20, 'Предел прочности': 474, 'Предел текучести': 300, 'Отн. удл.': 36, 'Отн. суж.': 61},
        {'Клеймо': '2-2', 'Температура': 20, 'Предел прочности': 466, 'Предел текучести': 290, 'Отн. удл.': 37, 'Отн. суж.': 63},
        {'Клеймо': '2-3', 'Температура': 403, 'Предел прочности': 443, 'Предел текучести': 264, 'Отн. удл.': 27, 'Отн. суж.': 65},
        {'Клеймо': '2-4', 'Температура': 403, 'Предел прочности': 444, 'Предел текучести': 305, 'Отн. удл.': 25, 'Отн. суж.': 62},
        {'Клеймо': '3-1', 'Температура': 20, 'Предел прочности': 488, 'Предел текучести': 301, 'Отн. удл.': 30, 'Отн. суж.': 60},
        {'Клеймо': '3-2', 'Температура': 20, 'Предел прочности': 487, 'Предел текучести': 305, 'Отн. удл.': 34, 'Отн. суж.': 60},
        {'Клеймо': '3-3', 'Температура': 403, 'Предел прочности': 428, 'Предел текучести': 250, 'Отн. удл.': 31, 'Отн. суж.': 65},
        {'Клеймо': '3-4', 'Температура': 403, 'Предел прочности': 427, 'Предел текучести': 249, 'Отн. удл.': 32, 'Отн. суж.': 63},
        {'Клеймо': '4-1', 'Температура': 20, 'Предел прочности': 525, 'Предел текучести': 401, 'Отн. удл.': 28, 'Отн. суж.': 59},
        {'Клеймо': '4-2', 'Температура': 20, 'Предел прочности': 520, 'Предел текучести': 336, 'Отн. удл.': 35, 'Отн. суж.': 60},
        {'Клеймо': '4-3', 'Температура': 403, 'Предел прочности': 450, 'Предел текучести': 242, 'Отн. удл.': 28, 'Отн. суж.': 60},
        {'Клеймо': '4-4', 'Температура': 403, 'Предел прочности': 447, 'Предел текучести': 246, 'Отн. удл.': 29, 'Отн. суж.': 62},
        {'Клеймо': '5-1', 'Температура': 20, 'Предел прочности': 494, 'Предел текучести': 266, 'Отн. удл.': 39, 'Отн. суж.': 60},
        {'Клеймо': '5-2', 'Температура': 20, 'Предел прочности': 496, 'Предел текучести': 273, 'Отн. удл.': 35, 'Отн. суж.': 59},
        {'Клеймо': '5-3', 'Температура': 403, 'Предел прочности': 430, 'Предел текучести': 232, 'Отн. удл.': 31, 'Отн. суж.': 64},
        {'Клеймо': '5-4', 'Температура': 403, 'Предел прочности': 436, 'Предел текучести': 224, 'Отн. удл.': 28, 'Отн. суж.': 68},
        {'Клеймо': '6-1', 'Температура': 20, 'Предел прочности': 502, 'Предел текучести': 295, 'Отн. удл.': 31, 'Отн. суж.': 59},
        {'Клеймо': '6-2', 'Температура': 20, 'Предел прочности': 503, 'Предел текучести': 294, 'Отн. удл.': 34, 'Отн. суж.': 55},
        {'Клеймо': '6-3', 'Температура': 403, 'Предел прочности': 469, 'Предел текучести': 254, 'Отн. удл.': 27, 'Отн. суж.': 64},
        {'Клеймо': '6-4', 'Температура': 403, 'Предел прочности': 454, 'Предел текучести': 223, 'Отн. удл.': 24, 'Отн. суж.': 65},
        {'Клеймо': '7-1', 'Температура': 20, 'Предел прочности': 504, 'Предел текучести': 329, 'Отн. удл.': 28, 'Отн. суж.': 58},
        {'Клеймо': '7-2', 'Температура': 20, 'Предел прочности': 499, 'Предел текучести': 314, 'Отн. удл.': 35, 'Отн. суж.': 57},
        {'Клеймо': '7-3', 'Температура': 403, 'Предел прочности': 459, 'Предел текучести': 278, 'Отн. удл.': 28, 'Отн. суж.': 67},
        {'Клеймо': '7-4', 'Температура': 403, 'Предел прочности': 457, 'Предел текучести': 264, 'Отн. удл.': 24, 'Отн. суж.': 63},
    ]
    
    return pd.DataFrame(test_data)

def create_detailed_dataframe(df, mapping=None, steel_grade='20'):
    """Создание детализированной таблицы с добавлением нормативных значений"""
    if df.empty:
        return pd.DataFrame(), []
    
    # Извлекаем номер трубы из клейма
    df['Номер трубы'] = df['Клеймо'].apply(lambda x: int(x.split('-')[0]) if '-' in str(x) else 0)
    df['Номер образца'] = df['Клеймо'].apply(lambda x: int(x.split('-')[1]) if '-' in str(x) else 0)
    
    # Определяем порядок следования образцов
    if mapping:
        sorted_pipes = []
        other_pipes = []
        
        for pipe_num in df['Номер трубы'].unique():
            if pipe_num in mapping:
                sorted_pipes.append(pipe_num)
            else:
                other_pipes.append(pipe_num)
        
        sorted_pipes.sort(key=lambda x: mapping[x]['order'])
        other_pipes.sort()
        ordered_pipes = sorted_pipes + other_pipes
        
        df['Порядок'] = df['Номер трубы'].apply(
            lambda x: mapping.get(x, {}).get('order', 999 + x))
        df['Новое название'] = df['Номер трубы'].apply(
            lambda x: mapping.get(x, {}).get('new_name', f"Труба {x}"))
        
        df = df.sort_values(['Порядок', 'Температура', 'Номер образца'])
    else:
        df['Новое название'] = df['Номер трубы'].apply(lambda x: f"Труба {x}")
        df = df.sort_values(['Номер трубы', 'Температура', 'Номер образца'])
        ordered_pipes = sorted(df['Номер трубы'].unique())
    
    detailed_rows = []
    non_conformities = []
    
    # Храним границы образцов для объединения ячеек в Word
    sample_boundaries = []
    
    # Проходим по трубам в нужном порядке
    for pipe_num in ordered_pipes:
        pipe_data = df[df['Номер трубы'] == pipe_num]
        
        if mapping and pipe_num in mapping:
            pipe_name = mapping[pipe_num]['new_name']
        else:
            pipe_name = f"Труба {pipe_num}"
        
        # Запоминаем начало образца
        start_index = len(detailed_rows)
        
        # Группируем по температуре
        for temp in sorted(pipe_data['Температура'].unique()):
            temp_data = pipe_data[pipe_data['Температура'] == temp]
            
            # Добавляем строки для каждого образца
            for _, row in temp_data.iterrows():
                row_data = {
                    'Образец': pipe_name,
                    'Температура, °C': temp,
                    'Предел прочности, МПа': int(round(row['Предел прочности'])),
                    'Предел текучести, МПа': int(round(row['Предел текучести'])),
                    'Отн. удл., %': int(round(row['Отн. удл.'])),
                    'Отн. суж., %': int(round(row['Отн. суж.']))
                }
                detailed_rows.append(row_data)
                
                # Проверяем на соответствие нормативам
                row_index = len(detailed_rows) - 1
                if temp <= 20:
                    if not check_against_normative(row_data['Предел прочности, МПа'], temp, 'strength', steel_grade):
                        non_conformities.append((row_index, 2))
                    if not check_against_normative(row_data['Предел текучести, МПа'], temp, 'yield', steel_grade):
                        non_conformities.append((row_index, 3))
                    if not check_against_normative(row_data['Отн. удл., %'], temp, 'elongation', steel_grade):
                        non_conformities.append((row_index, 4))
                    if not check_against_normative(row_data['Отн. суж., %'], temp, 'reduction', steel_grade):
                        non_conformities.append((row_index, 5))
                else:
                    if not check_against_normative(row_data['Предел текучести, МПа'], temp, 'yield', steel_grade, is_high_temp=True):
                        non_conformities.append((row_index, 3))
            
            # Добавляем строку со средними значениями (Среднее в столбце температуры)
            if len(temp_data) > 0:
                avg_row = {
                    'Образец': pipe_name,
                    'Температура, °C': 'Среднее',
                    'Предел прочности, МПа': int(round(temp_data['Предел прочности'].mean())),
                    'Предел текучести, МПа': int(round(temp_data['Предел текучести'].mean())),
                    'Отн. удл., %': int(round(temp_data['Отн. удл.'].mean())),
                    'Отн. суж., %': int(round(temp_data['Отн. суж.'].mean()))
                }
                detailed_rows.append(avg_row)
                
                # Проверяем средние значения на соответствие нормативам
                row_index = len(detailed_rows) - 1
                if temp <= 20:
                    if not check_against_normative(avg_row['Предел прочности, МПа'], temp, 'strength', steel_grade):
                        non_conformities.append((row_index, 2))
                    if not check_against_normative(avg_row['Предел текучести, МПа'], temp, 'yield', steel_grade):
                        non_conformities.append((row_index, 3))
                    if not check_against_normative(avg_row['Отн. удл., %'], temp, 'elongation', steel_grade):
                        non_conformities.append((row_index, 4))
                    if not check_against_normative(avg_row['Отн. суж., %'], temp, 'reduction', steel_grade):
                        non_conformities.append((row_index, 5))
                else:
                    if not check_against_normative(avg_row['Предел текучести, МПа'], temp, 'yield', steel_grade, is_high_temp=True):
                        non_conformities.append((row_index, 3))
        
        # Запоминаем конец образца
        end_index = len(detailed_rows) - 1
        sample_boundaries.append((start_index, end_index, pipe_name))
    
    # Удаляем последнюю пустую строку если есть
    if detailed_rows and all(v == '' for v in detailed_rows[-1].values()):
        detailed_rows.pop()
    
    # Добавляем нормативные значения
    steel_data = STEEL_GRADES.get(steel_grade, STEEL_GRADES['20'])
    
    normative_start = len(detailed_rows)
    detailed_rows.append({
        'Образец': f'Требования для {steel_data["name"]}',
        'Температура, °C': 20,
        'Предел прочности, МПа': f'{steel_data["room_temp"]["strength_range"][0]}-{steel_data["room_temp"]["strength_range"][1]}',
        'Предел текучести, МПа': f'не менее {steel_data["room_temp"]["yield_min"]}',
        'Отн. удл., %': f'не менее {steel_data["room_temp"]["elongation_min"]}',
        'Отн. суж., %': f'не менее {steel_data["room_temp"]["reduction_min"]}'
    })
    
    # Добавляем нормативные значения для повышенных температур
    unique_temps = sorted([t for t in df['Температура'].unique() if t > 20])
    
    for temp in unique_temps:
        normative_yield = get_interpolated_yield(steel_grade, temp)
        
        detailed_rows.append({
            'Образец': f'Требования для {steel_data["name"]}',
            'Температура, °C': temp,
            'Предел прочности, МПа': '-',
            'Предел текучести, МПа': f'не менее {normative_yield}',
            'Отн. удл., %': '-',
            'Отн. суж., %': '-'
        })
    
    detailed_df = pd.DataFrame(detailed_rows)
    return detailed_df, non_conformities, sample_boundaries

def create_summary_table(df, mapping=None, steel_grade='20'):
    """Создание сводной таблицы со средними пределами текучести при повышенной температуре"""
    if df.empty:
        return pd.DataFrame(), []
    
    # Извлекаем номер трубы
    df['Номер трубы'] = df['Клеймо'].apply(lambda x: int(x.split('-')[0]) if '-' in str(x) else 0)
    
    # Определяем порядок
    if mapping:
        summary_rows = []
        for pipe_num in df['Номер трубы'].unique():
            pipe_data = df[df['Номер трубы'] == pipe_num]
            high_temp_data = pipe_data[pipe_data['Температура'] > 20]
            
            if not high_temp_data.empty:
                avg_yield = int(round(high_temp_data['Предел текучести'].mean()))
                
                if pipe_num in mapping:
                    pipe_name = mapping[pipe_num]['new_name']
                    order = mapping[pipe_num]['order']
                else:
                    pipe_name = f"Труба {pipe_num}"
                    order = 999 + pipe_num
                
                summary_rows.append({
                    'Порядок': order,
                    'Образец': pipe_name,
                    'Средний предел текучести, МПа': avg_yield
                })
        
        # Сортируем по порядку (такому же как в основной таблице)
        summary_df = pd.DataFrame(summary_rows)
        if not summary_df.empty:
            summary_df = summary_df.sort_values('Порядок').drop('Порядок', axis=1)
    else:
        summary_rows = []
        for pipe_num in sorted(df['Номер трубы'].unique()):
            pipe_data = df[df['Номер трубы'] == pipe_num]
            high_temp_data = pipe_data[pipe_data['Температура'] > 20]
            
            if not high_temp_data.empty:
                avg_yield = int(round(high_temp_data['Предел текучести'].mean()))
                
                summary_rows.append({
                    'Образец': f"Труба {pipe_num}",
                    'Средний предел текучести, МПа': avg_yield
                })
        
        summary_df = pd.DataFrame(summary_rows)
    
    temperatures_above_20 = sorted([t for t in df['Температура'].unique() if t > 20])
    return summary_df, temperatures_above_20

def create_word_report(detailed_df, summary_df, high_temps, non_conformities, sample_boundaries, steel_grade='20'):
    """Создание Word документа с таблицами"""
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Заголовок
    title = doc.add_paragraph('Таблица механических свойств')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].bold = True
    
    # Информация о марке стали
    steel_info = STEEL_GRADES.get(steel_grade, STEEL_GRADES['20'])
    steel_para = doc.add_paragraph(f'Марка стали для сравнения: {steel_info["name"]}')
    steel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")
    date_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Таблица 1
    doc.add_paragraph('1. Результаты механических испытаний образцов')
    doc.paragraphs[-1].runs[0].bold = True
    
    # Создаем таблицу
    if not detailed_df.empty:
        # Определяем, где начинаются нормативные значения (после всех образцов)
        normative_start = None
        for i, row in detailed_df.iterrows():
            if 'Требования' in str(row['Образец']):
                normative_start = i
                break
        
        table1 = doc.add_table(rows=len(detailed_df)+1, cols=len(detailed_df.columns))
        table1.style = 'Table Grid'
        table1.autofit = False
        
        # Заголовки
        headers = detailed_df.columns.tolist()
        for i, header in enumerate(headers):
            cell = table1.cell(0, i)
            cell.text = str(header)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Данные
        for i, row in detailed_df.iterrows():
            for j, col in enumerate(headers):
                cell = table1.cell(i+1, j)
                value = str(row[col]) if pd.notna(row[col]) else ''
                cell.text = value
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # Жирный шрифт для средних значений и нормативных строк
                if 'Среднее' in value or 'Требования' in str(row.get('Образец', '')):
                    cell.paragraphs[0].runs[0].font.bold = True
                
                # Выделение красным для несоответствий (только для строк с образцами, не для нормативных)
                if (i, j) in non_conformities and (normative_start is None or i < normative_start):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Объединение ячеек для названий образцов
        for start_idx, end_idx, pipe_name in sample_boundaries:
            if start_idx <= end_idx:
                # Объединяем ячейки в первом столбце от start_idx+1 до end_idx+1
                # (+1 потому что первая строка - заголовки)
                start_cell = table1.cell(start_idx + 1, 0)
                end_cell = table1.cell(end_idx + 1, 0)
                start_cell.merge(end_cell)
                
                # Центрируем текст по вертикали и горизонтали
                start_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Таблица 2
    if not summary_df.empty:
        if high_temps:
            temp_str = ", ".join(map(str, high_temps))
            title2 = doc.add_paragraph(f'2. Средние пределы текучести при повышенной температуре ({temp_str}°C)')
        else:
            title2 = doc.add_paragraph('2. Средние пределы текучести при повышенной температуре')
        title2.runs[0].bold = True
        
        table2 = doc.add_table(rows=len(summary_df)+1, cols=len(summary_df.columns))
        table2.style = 'Table Grid'
        
        # Заголовки
        headers2 = summary_df.columns.tolist()
        for i, header in enumerate(headers2):
            cell = table2.cell(0, i)
            cell.text = str(header)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Данные
        for i, row in summary_df.iterrows():
            for j, col in enumerate(headers2):
                cell = table2.cell(i+1, j)
                value = str(row[col]) if pd.notna(row[col]) else ''
                cell.text = value
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Сохраняем в BytesIO
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def main():
    """Основная функция"""
    st.markdown('<h1 class="main-header">📊 Обработчик протоколов механических испытаний</h1>', unsafe_allow_html=True)
    
    # Информационный блок
    st.markdown("""
    <div class="info-box">
    <h4>📁 Загрузите файлы для обработки</h4>
    <p>1. Протокол испытаний (DOCX) - обязательный<br>
    2. Файл соответствия названий (Excel) - опционально, для переименования образцов</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Два загрузчика файлов
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Протокол испытаний")
        uploaded_protocol = st.file_uploader(
            "Загрузите протокол испытаний (DOCX)",
            type=['docx'],
            key="protocol",
            help="Основной файл с результатами механических испытаний"
        )
    
    with col2:
        st.subheader("📊 Файл соответствия")
        uploaded_mapping = st.file_uploader(
            "Загрузите файл соответствия названий (Excel)",
            type=['xlsx', 'xls'],
            key="mapping",
            help="Excel файл с двумя столбцами: новое название и номер из протокола"
        )
    
    # Боковая панель
    with st.sidebar:
        st.header("⚙️ Настройки")
        st.markdown("---")
        
        st.subheader("Параметры обработки")
        use_test_data = st.checkbox("Использовать тестовые данные", value=True,
                                   help="Использовать примерные данные для демонстрации")
        
        st.subheader("Выбор марки стали")
        steel_grade = st.selectbox(
            "Выберите марку стали для сравнения:",
            options=list(STEEL_GRADES.keys()),
            format_func=lambda x: STEEL_GRADES[x]['name'],
            index=0
        )
        
        steel_info = STEEL_GRADES[steel_grade]
        with st.expander(f"📋 Нормативные значения для {steel_info['name']}"):
            st.write(f"**Описание:** {steel_info['description']}")
            st.write("**При 20°C:**")
            st.write(f"- Предел прочности: {steel_info['room_temp']['strength_range'][0]}-{steel_info['room_temp']['strength_range'][1]} МПа")
            st.write(f"- Предел текучести: не менее {steel_info['room_temp']['yield_min']} МПа")
            st.write(f"- Относительное удлинение: не менее {steel_info['room_temp']['elongation_min']}%")
            st.write(f"- Относительное сужение: не менее {steel_info['room_temp']['reduction_min']}%")
            
            if steel_info['high_temp_points']:
                st.write("**При повышенных температурах:**")
                for temp, value in sorted(steel_info['high_temp_points']):
                    st.write(f"- {temp}°C: предел текучести не менее {value} МПа")
    
    # Обработка файлов
    if uploaded_protocol is not None or use_test_data:
        try:
            with st.spinner("📊 Обработка данных..."):
                # Парсим файл соответствия если есть
                mapping = {}
                if uploaded_mapping is not None:
                    mapping = parse_mapping_file(uploaded_mapping)
                    if mapping:
                        st.success(f"✅ Загружено {len(mapping)} соответствий названий")
                
                # Получаем данные протокола
                if use_test_data:
                    df = get_test_data()
                    file_source = "тестовые данные"
                else:
                    file_content = uploaded_protocol.read()
                    df = parse_protocol_from_docx(file_content)
                    file_source = uploaded_protocol.name
                
                if df.empty:
                    st.error("Не удалось извлечь данные из файла.")
                    st.info("Попробуйте включить опцию 'Использовать тестовые данные'")
                    return
                
                # Создаем таблицы
                detailed_df, non_conformities, sample_boundaries = create_detailed_dataframe(df, mapping, steel_grade)
                summary_df, high_temps = create_summary_table(df, mapping, steel_grade)
                
                # Показываем статистику
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Обработано образцов", len(df))
                with col2:
                    unique_pipes = df['Клеймо'].apply(lambda x: str(x).split('-')[0] if '-' in str(x) else '0').nunique()
                    st.metric("Количество труб", unique_pipes)
                with col3:
                    temps = sorted(df['Температура'].unique())
                    st.metric("Температуры испытаний", f"{len(temps)} видов")
                
                # Показываем информацию о несоответствиях
                if non_conformities:
                    st.warning(f"⚠️ Найдены {len(non_conformities)} значений, не соответствующих нормативам")
                
                # Предпросмотр
                st.subheader("📋 Предпросмотр основной таблицы")
                st.dataframe(detailed_df, use_container_width=True, hide_index=True)
                
                # Показываем информацию о структуре таблицы в Word
                with st.expander("📝 Структура Word отчета"):
                    st.write("**Особенности форматирования в Word:**")
                    st.write("1. Название образца объединено в одну ячейку для всех его строк")
                    st.write("2. Убраны лабораторные клейма образцов")
                    st.write("3. В столбце температуры для средних значений указано 'Среднее'")
                    st.write("4. Несоответствующие значения выделены красным цветом")
                    st.write("5. Нормативные значения добавлены в конец таблицы")
                
                if not summary_df.empty:
                    st.subheader("📊 Предпросмотр сводной таблицы")
                    st.dataframe(summary_df, use_container_width=True, hide_index=True)
                    
                    # Показываем порядок образцов в сводной таблице
                    with st.expander("📋 Порядок образцов в сводной таблице"):
                        st.write("Образцы отсортированы в том же порядке, что и в основной таблице:")
                        for i, row in summary_df.iterrows():
                            st.write(f"{i+1}. {row['Образец']}: {row['Средний предел текучести, МПа']} МПа")
                
                # Создание Word документа
                st.subheader("📥 Скачать отчет")
                
                doc_bytes = create_word_report(detailed_df, summary_df, high_temps, non_conformities, sample_boundaries, steel_grade)
                
                # Кнопка скачивания
                filename = f"Таблица_механических_свойств_{steel_grade}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                
                st.download_button(
                    label=f"⬇️ Скачать отчет в Word ({STEEL_GRADES[steel_grade]['name']})",
                    data=doc_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
        except Exception as e:
            st.error(f"Ошибка при обработке: {str(e)}")
    
    else:
        # Инструкция
        st.info("👈 Загрузите протокол испытаний (DOCX файл) для начала обработки")

if __name__ == "__main__":
    main()
